"""Generate PRD v0.2 (Word) for AI与智能体服务管理平台.

Run: python gen_prd_v0_2.py
Output: AI与智能体服务管理平台_PRD_v0.2.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = r"F:\agentOperation\AI与智能体服务管理平台_PRD_v0.2.docx"


def set_cn_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 16, 2: 14, 3: 12.5, 4: 11.5}
    colors = {0: (15, 76, 129), 1: (15, 76, 129), 2: (15, 95, 174), 3: (19, 95, 174), 4: (60, 60, 60)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level > 0 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    set_cn_font(r, size=sizes.get(level, 11), bold=True, color=colors.get(level, (0, 0, 0)))
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_para(doc, text, size=10.5, bold=False, color=None, indent=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_cn_font(r, size=size)
    return p


def add_kv_bullet(doc, key, value, level=0, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{key}：")
    set_cn_font(r1, size=size, bold=True, color=(15, 76, 129))
    r2 = p.add_run(value)
    set_cn_font(r2, size=size)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color=(15, 76, 129)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_cn_font(r, size=10, bold=True, color=(255, 255, 255))
        # 背景色
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), f"{header_color[0]:02X}{header_color[1]:02X}{header_color[2]:02X}")
        tc_pr.append(shd)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 数据行
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.paragraph_format.line_spacing = 1.3
            r = p.add_run(str(val))
            set_cn_font(r, size=9.5)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    # 空行
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, text, kind="info"):
    """信息/告警块。kind: info(蓝)/warn(橙)/danger(红)/success(绿)"""
    colors = {
        "info": (220, 235, 252),
        "warn": (254, 243, 199),
        "danger": (254, 226, 226),
        "success": (220, 252, 231),
    }
    border_colors = {
        "info": (59, 130, 246),
        "warn": (245, 158, 11),
        "danger": (239, 68, 68),
        "success": (34, 197, 94),
    }
    fill = colors.get(kind, colors["info"])
    bd = border_colors.get(kind, border_colors["info"])
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    set_cn_font(r, size=10, bold=False, color=(60, 60, 60))
    # 段落背景
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), f"{fill[0]:02X}{fill[1]:02X}{fill[2]:02X}")
    pPr.append(shd)
    # 左侧边框
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), "E5E7EB")
        pBdr.append(e)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), f"{bd[0]:02X}{bd[1]:02X}{bd[2]:02X}")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================== 开始构建文档 ==============================
doc = Document()

# 页边距
for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4)
    s.right_margin = Cm(2.4)

# 默认正文样式
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.font.size = Pt(10.5)

# ===== 封面 =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(80)
title.paragraph_format.space_after = Pt(8)
r = title.add_run("AI与智能体服务管理平台")
set_cn_font(r, size=28, bold=True, color=(15, 76, 129))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(50)
r = sub.add_run("产品需求文档（PRD） v0.2")
set_cn_font(r, size=16, bold=False, color=(60, 60, 60))

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("—— 基于v0.1 Demo评审问题的改进版本 ——")
set_cn_font(r, size=11, color=(120, 120, 120))

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.paragraph_format.space_before = Pt(60)
for line in ["文档版本：v0.2（初稿，待评审）", "编制日期：2026-07-20", "文档状态：需求梳理中", "适用范围：江苏省医保局 AI与智能体服务管理平台"]:
    r = info_p.add_run(line + "\n")
    set_cn_font(r, size=11, color=(80, 80, 80))

doc.add_page_break()

# ===== 0. 文档说明 =====
add_heading(doc, "0  文档说明", level=1)

add_heading(doc, "0.1 编制目的", level=2)
add_para(doc, "本文档基于v0.1 Demo评审反馈，对AI与智能体服务管理平台的需求进行系统性补全与修订，重点解决Demo中存在的功能缺失、流程断点、角色与权限模型不完整、数据闭环未打通等问题，作为后续详细设计与开发实施的输入。本文档不替代最终设计文档，仅用于在需求梳理工具中进一步细化。", indent=True)

add_heading(doc, "0.2 修订记录", level=2)
add_table(doc,
    ["版本", "日期", "修订内容", "编制/修订人"],
    [
        ["v0.1", "2026-07-16", "基于建设内容文档形成首版PRD与Demo界面骨架", "—"],
        ["v0.2", "2026-07-20", "针对Demo评审问题，补全缺失模块、修正流程断点、明确角色权限、增加试点与多主体协同、增加端到端数据流转闭环", "—"],
    ],
    col_widths=[1.5, 2.5, 9, 2])

add_heading(doc, "0.3 术语与缩写", level=2)
add_table(doc,
    ["术语", "说明"],
    [
        ["AI服务", "在本平台中指可被订阅、调用的模型、智能体及其组合，统一以服务形式纳管"],
        ["词元（Token）", "AI服务调用的统一计量单位，包括但不限于LLM Token、检查例次、调用次数等"],
        ["五类分类", "通用基础大模型 / 医保自研专属大模型 / 医保基金监管共建模型 / 省头部医疗机构共建垂直模型 / 市场化合规生态AI产品"],
        ["三级准入", "高风险（资质核验+技术测评+临床验收）/ 中风险（资质核验+简化技术测试）/ 低风险（备案上架）"],
        ["多主体", "医保部门、医疗机构、运营单位、AI企业、科研机构、医院自研团队等参与方"],
        ["试点", "在特定城市、医院、项目范围内进行的限定性试运行，配套独立授权与计量规则"],
        ["RABC", "基于角色的访问控制（Role-Based Access Control）"],
    ],
    col_widths=[3, 12])

doc.add_page_break()

# ===== 1. 项目背景与建设目标 =====
add_heading(doc, "1  项目背景与建设目标", level=1)

add_heading(doc, "1.1 项目背景", level=2)
add_para(doc, "江苏省医保局在AI与智能体服务应用方面已积累一定基础，但仍面临服务资源分散、准入口径不统一、计量计费不闭环、多主体协同不畅、试点成效难以量化等问题。本项目旨在建设统一的AI与智能体服务管理平台，将各类模型、智能体、数据资源、知识资源、平台工具及相关供给主体纳入统一管理，形成全省AI服务资源底账与运营中枢。", indent=True)

add_heading(doc, "1.2 建设目标", level=2)
add_kv_bullet(doc, "服务资源统一纳管", "形成全省AI服务资源底账，统一服务目录、资源底账、准入状态和使用范围口径，减少重复建设与资源闲置。")
add_kv_bullet(doc, "准入运营闭环管理", "建立贯穿 准入-上架-运行-评价-退出 的全生命周期管理，过程可留痕、可核验、可追溯。")
add_kv_bullet(doc, "词元计量计费闭环", "从概念设计转向实际运行管理，统一计量规则、对账规则与异常处置规则，支撑使用可计量、成本可核算、结果可评价。")
add_kv_bullet(doc, "多主体协同治理", "明确医保部门、医疗机构、运营单位、AI企业、科研机构、医院自研团队等主体的接入边界、授权、权限与责任留痕。")
add_kv_bullet(doc, "试点成效可量化", "支持试点城市、试点医院、专项项目的独立授权、独立计量与独立成效评估，为推广决策提供依据。")

add_heading(doc, "1.3 用户角色概览", level=2)
add_para(doc, "平台按 前台门户 + 后台管理 两大板块组织，识别以下角色（v0.1仅支持3种，v0.2补全）：", indent=True)
add_table(doc,
    ["角色", "归属板块", "主要职责", "v0.1是否实现"],
    [
        ["访客", "前台门户", "浏览门户、服务市场，无需登录", "部分（公开访问）"],
        ["机构用户", "机构工作台", "本院AI服务订阅、配额管控、调用权限申请、费用结算、评价反馈", "部分（缺调用权限申请、评价反馈）"],
        ["开发者", "开发者中心", "上架申请、测试协同、版本迭代、收益对账", "部分（缺版本迭代）"],
        ["平台运营", "管理中心", "运营管理、运营分析、对账、计费", "部分（流程不完整）"],
        ["平台审核", "管理中心", "准入审核、资质审核、整改跟踪", "未独立区分（混入运营）"],
        ["平台管理员", "管理中心", "系统配置、用户与角色、接口与集成", "部分"],
        ["医保监管", "管理中心", "基金监管共建模型监管、场景应用分析", "未实现"],
    ],
    col_widths=[2.5, 2.5, 7.5, 2.5])

add_note(doc, "改进点①：v0.1将 运营/审核/管理 合并为admin单角色，导致权限边界模糊、操作无法精细留痕。v0.2按职责拆分为平台运营、平台审核、平台管理员三类，医保监管作为独立只读角色。", "warn")

doc.add_page_break()

# ===== 2. v0.1 Demo问题与改进点 =====
add_heading(doc, "2  v0.1 Demo问题盘点与改进方向", level=1)
add_para(doc, "本章列出Demo评审中识别的主要缺失与逻辑不合理点，作为v0.2需求补全的依据。每条问题均标注影响范围与改进方向，后续章节将给出具体功能要求。", indent=True)

add_heading(doc, "2.1 角色与权限模型", level=2)
add_table(doc,
    ["问题编号", "问题描述", "影响", "改进方向"],
    [
        ["P-01", "仅org/developer/admin三种角色，审核、运营、管理职责未分离", "权限边界模糊、操作无法精细留痕", "拆分为运营、审核、管理员三类，新增医保监管只读角色"],
        ["P-02", "登录为mock点击式，无真实身份认证", "无法对接统一身份认证", "预留SSO/OAuth接口，登录页支持账号密码+SSO双通道"],
        ["P-03", "无角色权限矩阵，仅按路由做粗粒度拦截", "越权风险", "输出RABC权限矩阵，菜单/按钮/数据三级权限"],
        ["P-04", "无访客概念，公开页与登录页跳转混乱", "用户体验差", "明确访客可访问范围，未登录访问受限页时引导登录"],
    ],
    col_widths=[1.8, 6.5, 3, 4.2])

add_heading(doc, "2.2 业务流程与闭环", level=2)
add_table(doc,
    ["问题编号", "问题描述", "影响", "改进方向"],
    [
        ["P-05", "服务订阅申请无审批环节，订阅列表只是数据展示", "无法体现 申请-审批-生效 闭环", "增加订阅申请单、审批工作台、状态机"],
        ["P-06", "机构 模型调用权限申请 页面缺失（PRD 2.3.3）", "机构无法为院内主体申请调用权限", "新增调用权限申请页，与订阅、配额联动"],
        ["P-07", "机构 服务评价反馈 页面缺失（PRD 2.3.3）", "评价数据无采集入口，质量评价无来源", "新增评价提交页，纳入服务质量评价体系"],
        ["P-08", "开发者 版本迭代管理 页面缺失（PRD 2.4.3）", "版本变更无管理，复审触发规则不明", "新增版本迭代页，按风险等级触发复审"],
        ["P-09", "开发者上架申请与平台准入审核之间无显式流转", "两侧看板割裂", "上架申请单作为准入审核的输入，状态双向同步"],
        ["P-10", "退订/退款流程缺失", "订阅到期或撤销无处理路径", "新增退订申请单、剩余词元结算规则"],
        ["P-11", "资质到期管理缺失", "资质失效后服务仍在售风险", "新增到期提醒、续证流程、自动下架规则"],
        ["P-12", "异常对账只识别不处置", "异常挂账无闭环", "增加核验、退费、追责处置流程与状态机"],
    ],
    col_widths=[1.8, 6.5, 3, 4.2])

add_heading(doc, "2.3 模块功能完整性", level=2)
add_table(doc,
    ["问题编号", "问题描述", "改进方向"],
    [
        ["P-13", "资产总览管理仅按 数据/知识/工具 三类资源分页，缺少资产台账、资产分类总览、资产使用概览三个功能维度", "在三类资源之上增加资产总览首页，提供台账视图、分布视图、使用视图三个Tab"],
        ["P-14", "资产与AI服务的关联关系缺失（某模型用了哪些数据/知识/工具）", "增加资产-服务关联视图，模型/智能体详情页展示所依赖资产"],
        ["P-15", "词元计费管理将规则、额度、包管理混在一个页面，无版本/生效失效管理", "拆分为独立子页，规则支持版本化、生效失效时间、审批留痕"],
        ["P-16", "对账查询导出无审计留痕", "导出操作纳入安全审计，导出文件含水印与追溯信息"],
        ["P-17", "运行监测仅展示大屏，无下钻与异常处置", "增加异常下钻、告警工单、处置闭环"],
        ["P-18", "质量评价无评价方案、指标配置、结果明细", "增加评价指标库、评价方案管理、评价结果明细与趋势"],
        ["P-19", "安全审计仅日志列表，无审计规则、告警处置", "增加审计规则配置、告警阈值、告警工单"],
        ["P-20", "运营分析与运营管理边界不清，分析维度未下钻", "运营分析独立模块，三个分析维度分别支持机构/服务/场景下钻"],
        ["P-21", "服务详情页缺少词元计量规则、试点案例、评价评分、相关推荐", "按PRD 2.2.3补全详情字段与交互"],
        ["P-22", "门户首页缺少公告详情、搜索、个人中心", "增加公告详情页、全站搜索、个人中心"],
    ],
    col_widths=[1.8, 7.5, 6.2])

add_heading(doc, "2.4 端到端数据流转", level=2)
add_table(doc,
    ["问题编号", "问题描述", "改进方向"],
    [
        ["P-23", "服务上架—订阅—配额—调用—计费—对账—评价 七环节未打通，数据各自孤立", "定义统一主数据（服务ID、订阅ID、配额ID、调用ID、账单ID、评价ID），保证跨模块可追溯"],
        ["P-24", "机构侧配额与平台侧额度配置数据不同步", "配额变更通过事件总线同步，两侧展示一致"],
        ["P-25", "调用明细与对账账单的口径不一致风险", "统一计费规则版本号，账单与明细绑定规则版本"],
        ["P-26", "试点项目的独立计量与成效评估无支撑", "新增试点管理模块，试点项目作为独立授权与计量单元"],
    ],
    col_widths=[1.8, 7.5, 6.2])

add_heading(doc, "2.5 非功能性问题", level=2)
add_table(doc,
    ["问题编号", "问题描述", "改进方向"],
    [
        ["P-27", "无消息通知/站内信/待办聚合", "增加消息中心、待办工作台"],
        ["P-28", "无API开放能力、API Key管理、调用统计、限流配置", "新增API开放平台子模块"],
        ["P-29", "敏感数据（患者数据、医保数据）无脱敏/加密/权限隔离说明", "在系统管理中补充数据安全配置章节"],
        ["P-30", "无错误处理、404/403页、引导帮助", "增加全局错误页与新手引导"],
    ],
    col_widths=[1.8, 7.5, 6.2])

doc.add_page_break()

# ===== 3. 系统总体设计 =====
add_heading(doc, "3  系统总体设计", level=1)

add_heading(doc, "3.1 总体架构", level=2)
add_para(doc, "系统围绕 数据资源层、平台能力层、服务应用层 三层展开：", indent=True)
add_kv_bullet(doc, "数据资源层", "沉淀服务目录、模型与智能体资产、知识与工具资源、词元消耗、用户机构、开发者信息、评价结果、审计日志等管理数据。")
add_kv_bullet(doc, "平台能力层", "承载模型编目、资产总览、词元计费、对账查询、运营管理、运营分析、开发者管理、用户管理、试点管理、系统管理等核心能力。")
add_kv_bullet(doc, "服务应用层", "面向医保AI资源统筹、服务准入审核、词元运营管理、运行监管分析、服务评价退出和多主体协同管理等业务工作。")

add_heading(doc, "3.2 功能板块划分", level=2)
add_table(doc,
    ["板块", "模块", "主要能力", "v0.2状态"],
    [
        ["前台门户", "门户首页", "平台介绍、公告动态、快捷入口、开发者入驻、搜索、个人中心", "补全"],
        ["前台门户", "AI服务市场", "分类浏览、检索匹配、服务详情、订阅申请", "补全详情与订阅流程"],
        ["前台门户", "机构工作台", "订阅管理、配额管理、调用权限申请、费用结算、评价反馈、院内用户", "新增3个子页"],
        ["前台门户", "开发者中心", "上架申请、测试协同、版本迭代、收益对账", "新增版本迭代"],
        ["管理中心", "模型编目管理", "分类管理、信息管理、资质材料", "保持"],
        ["管理中心", "资产总览管理", "资产台账、分类总览、使用概览 + 数据/知识/工具三类资源子页 + 资产-服务关联视图", "新增总览与关联视图"],
        ["管理中心", "词元计费管理", "计量管理、额度配置、消耗监测、计费规则、词元包管理", "拆分子页、规则版本化"],
        ["管理中心", "对账查询管理", "调用明细、周期账单、异常对账、结果导出", "补全处置流程"],
        ["管理中心", "运营管理", "服务准入、上架下架、运行监测、质量评价、安全审计", "补全处置与配置"],
        ["管理中心", "运营分析", "服务热度、服务质量、场景应用分析", "支持下钻"],
        ["管理中心", "开发者管理", "资质审核、整改跟踪", "保持"],
        ["管理中心", "用户与权限管理", "机构用户、角色权限、服务授权、行为留痕", "新增权限矩阵"],
        ["管理中心", "试点与多主体协同", "试点项目管理、多主体接入边界、协同授权", "新增模块"],
        ["管理中心", "系统管理", "基础参数、接口与集成、API开放平台、系统日志、数据安全", "新增API与数据安全"],
    ],
    col_widths=[1.8, 2.7, 7.5, 2.5])

add_heading(doc, "3.3 角色权限矩阵（RABC）", level=2)
add_para(doc, "下表给出主要功能与角色的访问关系（C=创建 R=查询 U=修改 D=删除 A=审批）：", indent=True)
add_table(doc,
    ["功能模块", "访客", "机构用户", "开发者", "平台运营", "平台审核", "平台管理员", "医保监管"],
    [
        ["门户/服务市场", "R", "R", "R", "R", "R", "R", "R"],
        ["服务订阅", "-", "C/R/U", "-", "A/R", "R", "R", "R"],
        ["配额管理", "-", "R/U(本院)", "-", "C/R/U", "A", "R", "R"],
        ["调用权限申请", "-", "C/R", "-", "A/R", "R", "R", "R"],
        ["上架申请", "-", "-", "C/R/U", "R", "A", "R", "R"],
        ["准入审核", "-", "-", "R(本人)", "R", "A", "R", "R"],
        ["模型编目", "-", "R", "R", "R", "R", "C/R/U/D", "R"],
        ["资产总览", "-", "R", "R", "R", "R", "C/R/U/D", "R"],
        ["词元计费规则", "-", "R", "R", "C/R/U", "A", "R", "R"],
        ["对账查询", "-", "R(本院)", "R(本人)", "C/R/U", "A", "R", "R"],
        ["运营监测", "-", "R(本院)", "R(本人)", "R/U", "R", "R", "R"],
        ["质量评价", "-", "R(本院)", "R(本人)", "C/R/U", "A", "R", "R"],
        ["安全审计", "-", "-", "-", "R", "R", "R/U", "R"],
        ["运营分析", "-", "R(本院)", "R(本人)", "R", "R", "R", "R"],
        ["开发者管理", "-", "-", "R(本人)", "R", "A", "R/U", "R"],
        ["用户与权限", "-", "R(本人)", "R(本人)", "R", "R", "C/R/U/D", "R"],
        ["试点管理", "-", "R", "R", "C/R/U", "A", "R/U", "R"],
        ["系统管理", "-", "-", "-", "R", "R", "C/R/U/D", "-"],
    ],
    col_widths=[3, 1.4, 1.7, 1.5, 1.7, 1.7, 1.8, 1.7])

add_note(doc, "改进点②：上表为初版权限矩阵，需在需求梳理工具中进一步细化到按钮级（如导出、下架、调整额度等敏感操作）。医保监管为只读角色，重点访问运营分析、对账、安全审计相关数据。", "info")

add_heading(doc, "3.4 主数据与数据流转闭环", level=2)
add_para(doc, "为保证七环节可追溯，定义以下主数据ID与流转关系：", indent=True)
add_kv_bullet(doc, "服务ID（service_id）", "AI服务的全局唯一标识，从上架申请通过后生成。")
add_kv_bullet(doc, "版本ID（version_id）", "同一服务的版本号，每次版本迭代新增，绑定计费规则版本。")
add_kv_bullet(doc, "订阅ID（subscription_id）", "机构订阅服务后生成，关联服务ID、机构ID、有效期、配额ID。")
add_kv_bullet(doc, "配额ID（quota_id）", "订阅生效后生成，关联订阅ID，记录词元额度、已用、可用。")
add_kv_bullet(doc, "调用权限ID（permission_id）", "机构为院内主体申请的调用权限，关联订阅ID与主体ID。")
add_kv_bullet(doc, "调用ID（call_id）", "每次AI服务调用生成，关联服务版本ID、权限ID、调用主体、词元消耗、计费规则版本。")
add_kv_bullet(doc, "账单ID（bill_id）", "按周期归集调用明细生成，关联机构ID、服务ID、周期、调用ID列表。")
add_kv_bullet(doc, "评价ID（evaluation_id）", "机构用户提交评价后生成，关联订阅ID与服务ID。")
add_kv_bullet(doc, "审计ID（audit_id）", "关键行为留痕，关联操作人、操作对象、操作类型。")

add_para(doc, "流转关系：上架申请单 → 服务ID+版本ID → 订阅申请单 → 订阅ID+配额ID → 调用权限申请 → 权限ID → 服务调用 → 调用ID → 周期账单 → 账单ID → 异常对账/评价/审计。任一环节均可向上向下追溯。", indent=True)

doc.add_page_break()

# ===== 4. 功能详述 - 前台门户 =====
add_heading(doc, "4  功能详述——前台门户", level=1)

# 4.1 门户首页
add_heading(doc, "4.1 门户首页", level=2)
add_para(doc, "面向所有访客与登录用户，作为平台统一对外展示与入口引导窗口。", indent=True)

add_heading(doc, "4.1.1 平台介绍展示", level=3)
add_bullet(doc, "集中展示平台定位、核心能力、服务范围、建设成效、各方主体等基本信息。")
add_bullet(doc, "支持内容动态更新（运营在后台维护，无需开发介入）。")

add_heading(doc, "4.1.2 公告动态发布", level=3)
add_bullet(doc, "发布平台通知、政策更新、模型上下架、系统维护、试点进展等管理公告。")
add_bullet(doc, "按类型分类管理：普通通知 / 政策更新 / 紧急公告 / 试点进展 / 上下架通知。")
add_bullet(doc, "关键管理事项按业务节点及时触达相关主体（站内信+邮件/短信）。")
add_bullet(doc, "【新增】公告详情页：支持富文本、附件、置顶、已读未读统计。")

add_heading(doc, "4.1.3 快捷入口导航", level=3)
add_bullet(doc, "提供AI服务市场、机构工作台、开发者中心、运营管理后台等业务门户统一入口。")
add_bullet(doc, "按用户角色智能推荐常用功能（登录后展示）。")
add_bullet(doc, "【新增】未登录用户展示 申请入驻 和 进入服务市场 两个主入口。")

add_heading(doc, "4.1.4 开发者入驻", level=3)
add_bullet(doc, "供给主体（AI企业、科研机构、医疗机构自研团队）注册入口。")
add_bullet(doc, "提交企业资质、医疗器械注册证或备案凭证、算法备案等材料。")
add_bullet(doc, "资质核验通过后获得平台接入资格。")
add_bullet(doc, "【新增】资质到期自动提醒续证（提前30/15/7天三级提醒）。")
add_bullet(doc, "【新增】入驻进度查询：提交→初审→复审→通过/驳回，全流程可视。")

add_heading(doc, "4.1.5 全站搜索【新增】", level=3)
add_bullet(doc, "支持按关键词搜索服务、开发者、公告、数据资源、知识、工具。")
add_bullet(doc, "搜索结果按类型分组展示，支持二次筛选与排序。")

add_heading(doc, "4.1.6 个人中心【新增】", level=3)
add_bullet(doc, "个人信息维护、密码修改、绑定手机/邮箱。")
add_bullet(doc, "我的消息、我的待办、我的申请、我的订阅（按角色展示）。")
add_bullet(doc, "登录历史、操作记录查询。")

# 4.2 AI服务市场
add_heading(doc, "4.2 AI服务市场", level=2)
add_para(doc, "平台服务货架，集中展示经统一评审准入后上架的AI模型与智能体，面向机构用户和开发者提供市场化服务能力。", indent=True)

add_heading(doc, "4.2.1 服务分类浏览", level=3)
add_bullet(doc, "按五类分类、应用场景、服务对象、开发者等维度组织已上架AI服务的统一展示。")
add_bullet(doc, "不同类别适用差异化的准入管控要求和词元计量规则（卡片上以标签提示）。")
add_bullet(doc, "提供最热/最新服务榜单。")

add_heading(doc, "4.2.2 服务检索匹配", level=3)
add_bullet(doc, "支持按关键词、供给方、能力类型、适用场景、风险等级等条件组合检索。")
add_bullet(doc, "检索结果支持多维度排序（热度、上线时间、评分等）。")
add_bullet(doc, "【新增】支持保存常用筛选组合为 我的视图。")

add_heading(doc, "4.2.3 服务详情展示", level=3)
add_para(doc, "形成完整服务档案视图，支撑服务适配性判断和选购决策。包含以下字段：", indent=True)
add_table(doc,
    ["字段分组", "字段", "说明"],
    [
        ["基础信息", "服务名称/简介/图标/分类/风险等级/开发者/上架时间/当前版本", "—"],
        ["能力说明", "能力描述/适用场景/能力边界/输入输出要求/示例", "—"],
        ["计量规则", "计量方式（Token/例次/调用次数）/计费规则版本/单价/试用额度", "【补全】v0.1缺失"],
        ["资质依据", "医疗器械注册证/算法备案/测试报告/合作材料", "—"],
        ["试点案例", "试点医院/试点场景/试点成效/案例说明", "【补全】v0.1缺失"],
        ["评价评分", "综合评分/分项评分/评价列表/好评率", "【补全】v0.1缺失"],
        ["相关推荐", "同类服务/同场景服务/同开发者服务", "【补全】v0.1缺失"],
        ["订阅入口", "订阅申请按钮/试用申请按钮/联系开发者", "【补全】订阅流程入口"],
    ],
    col_widths=[2.5, 6.5, 4.5])

add_heading(doc, "4.2.4 服务订阅申请", level=3)
add_bullet(doc, "机构用户在AI服务市场对所需AI服务发起订阅申请。")
add_bullet(doc, "申请单字段：订阅服务、院内使用范围、使用主体、额度预算、预期场景、预算来源（信息化经费/科室专项/科研经费）、有效期。")
add_bullet(doc, "【新增】订阅申请单状态机：草稿 → 已提交 → 平台运营初审 → 平台审核复审 → 通过/驳回 → 生效/失败。")
add_bullet(doc, "【新增】审批通过后生成订阅ID与配额ID，机构可在工作台查看。")
add_bullet(doc, "【新增】订阅试用：低风险服务支持7天免费试用，试用期内词元额度独立。")

# 4.3 机构工作台
add_heading(doc, "4.3 机构工作台", level=2)
add_para(doc, "面向医疗机构信息管理员，提供机构侧AI服务闭环管理能力。", indent=True)

add_heading(doc, "4.3.1 服务订阅管理", level=3)
add_bullet(doc, "集中查看本机构已订阅的AI服务清单、订阅状态、有效期、使用范围。")
add_bullet(doc, "支持续订、退订申请。")
add_bullet(doc, "【新增】退订流程：退订申请 → 剩余词元结算（按规则退还/不退还） → 平台审核 → 生效。")
add_bullet(doc, "订阅过程对接机构信息化经费、科室专项经费等预算来源。")

add_heading(doc, "4.3.2 用量配额管理", level=3)
add_bullet(doc, "对本机构订阅的AI服务进行词元额度分配、调整和管控。")
add_bullet(doc, "支持按科室、项目、场景、时间段等维度分配额度。")
add_bullet(doc, "配额规则与院内用户授权联动；超配额使用纳入预警和拦截。")
add_bullet(doc, "【新增】配额调整操作需二级审批（科室负责人+机构信息管理员）。")
add_bullet(doc, "【新增】配额预警阈值可配置（默认80%预警、95%拦截、100%冻结）。")

add_heading(doc, "4.3.3 模型调用权限申请【补全】", level=3)
add_para(doc, "v0.1缺失此页面，v0.2补全。", indent=True)
add_bullet(doc, "机构针对已订阅AI服务，为院内具体使用主体（科室、人员、系统）申请模型调用权限。")
add_bullet(doc, "申请字段：订阅服务、调用主体（科室/人员/系统）、调用范围、使用场景、额度约束、有效期。")
add_bullet(doc, "权限申请经审核后生效，生成调用权限ID。")
add_bullet(doc, "支持权限回收、权限变更、权限续期。")

add_heading(doc, "4.3.4 费用结算查询", level=3)
add_bullet(doc, "查看本机构词元消耗明细、周期费用账单和对账记录。")
add_bullet(doc, "支持按周期、服务、科室、项目等维度查询和核对。")
add_bullet(doc, "【新增】异常消耗可发起核验申请，关联平台异常对账工单。")
add_bullet(doc, "【新增】账单导出（PDF/Excel），含水印与追溯信息。")

add_heading(doc, "4.3.5 服务评价反馈【补全】", level=3)
add_para(doc, "v0.1缺失此页面，v0.2补全。评价数据作为平台服务质量评价的来源。", indent=True)
add_bullet(doc, "机构用户针对已订阅AI服务的使用体验、响应准确性、稳定性和业务适配性进行评价。")
add_bullet(doc, "评价字段：综合评分（1-5星）+ 分项评分（准确性/稳定性/时效性/适配性）+ 文字评价 + 截图附件。")
add_bullet(doc, "评价状态：待评价 / 已评价 / 已回复。开发者可对评价进行回复。")
add_bullet(doc, "评价纳入平台服务质量评价体系，作为服务续用、整改或下架的参考依据。")

add_heading(doc, "4.3.6 院内用户管理【新增】", level=3)
add_bullet(doc, "管理机构内使用AI服务的科室、人员账号。")
add_bullet(doc, "支持院内用户的调用权限分配、额度分配、使用记录查询。")
add_bullet(doc, "对接院内HIS/HRP系统的人员组织数据（接口对接）。")

# 4.4 开发者中心
add_heading(doc, "4.4 开发者中心", level=2)
add_para(doc, "面向已注册的AI产品供给方，提供一站式供给侧管理能力。", indent=True)

add_heading(doc, "4.4.1 上架申请管理", level=3)
add_bullet(doc, "引导开发者按统一口径提交服务说明、适用场景、能力边界、输入输出要求、测试材料、计量建议和部署要求等。")
add_bullet(doc, "形成完整上架申报档案。")
add_bullet(doc, "【新增】上架申请单与平台准入审核双向同步，开发者可查看审核进度与各环节意见。")
add_bullet(doc, "审核通过后完成服务发布；上架状态变更同步至AI服务市场前端展示。")

add_heading(doc, "4.4.2 测试协同管理", level=3)
add_bullet(doc, "开发者配合平台开展技术安全测评和临床场景落地验收。")
add_bullet(doc, "提交测试数据、测试环境和配合人员等资源。")
add_bullet(doc, "测试过程和结果在开发者中心统一查看，测试发现的问题纳入整改闭环。")
add_bullet(doc, "【新增】测试任务状态：待提交 → 测试中 → 测试通过/不通过 → 整改 → 复测。")

add_heading(doc, "4.4.3 版本迭代管理【补全】", level=3)
add_para(doc, "v0.1缺失此页面，v0.2补全。", indent=True)
add_bullet(doc, "已上架服务的版本更新、能力变更和模型迭代纳入统一管理。")
add_bullet(doc, "开发者提交版本变更说明和变更影响评估。")
add_bullet(doc, "【新增】版本迭代按风险等级决定是否触发复审：")
add_bullet(doc, "  - 低风险变更（UI/文案/性能优化）：仅备案，无需复审。", level=1)
add_bullet(doc, "  - 中风险变更（参数调整/能力扩展）：触发简化复审。", level=1)
add_bullet(doc, "  - 高风险变更（模型重训/核心算法变更/数据集变更）：触发全流程复审。", level=1)
add_bullet(doc, "【新增】版本灰度发布：支持按机构/试点范围灰度，灰度期结束自动全量。")
add_bullet(doc, "【新增】版本回滚：发现严重问题可一键回滚至上一稳定版本。")

add_heading(doc, "4.4.4 收益对账查询", level=3)
add_bullet(doc, "查看开发者所提供服务的调用词元消耗、收益分成、结算记录和到账情况。")
add_bullet(doc, "依据服务调用词元消耗获取对应收益。")
add_bullet(doc, "收益分配规则透明可查；对账结果可导出用于企业内部核算。")
add_bullet(doc, "【新增】收益争议申诉：开发者对账单有异议可发起申诉，进入平台运营处置流程。")

doc.add_page_break()

# ===== 5. 功能详述 - 管理中心 =====
add_heading(doc, "5  功能详述——管理中心", level=1)

# 5.1 模型编目管理
add_heading(doc, "5.1 模型编目管理", level=2)
add_heading(doc, "5.1.1 模型分类管理", level=3)
add_bullet(doc, "围绕基础大模型、专用模型、智能体等服务形态，形成统一的AI服务分类管理体系。")
add_bullet(doc, "覆盖服务来源、应用场景、服务对象、成熟程度和准入状态等管理口径。")
add_bullet(doc, "五类分类适用差异化的准入管控要求和词元计量规则。")

add_heading(doc, "5.1.2 模型信息管理", level=3)
add_bullet(doc, "以模型和智能体服务业务档案为基础，完整记录服务名称、供给单位、适用场景、能力边界、输入输出要求、使用条件和责任主体等信息。")

add_heading(doc, "5.1.3 资质材料管理", level=3)
add_bullet(doc, "服务准入所需的证照、备案、测试报告、合作材料等实行集中归集和核验。")
add_bullet(doc, "对临床辅助、疾病筛查、处方审核等敏感场景服务，形成上线前资质核验、合规判断和风险识别依据。")
add_bullet(doc, "【新增】资质到期管理：到期前30/15/7天三级提醒，到期未续证自动下架。")

# 5.2 资产总览管理
add_heading(doc, "5.2 资产总览管理", level=2)
add_para(doc, "形成全省AI服务资源底账，将模型、智能体、数据资源、知识资源和平台工具等统一纳入可视化管理范围。v0.1仅有三类资源子页，v0.2补全总览与关联视图。", indent=True)

add_heading(doc, "5.2.1 资产总览首页【新增】", level=3)
add_bullet(doc, "三个Tab：资产台账 / 资产分类总览 / 资产使用概览。")
add_bullet(doc, "资产台账：列出全部资产（模型/智能体/数据/知识/工具），记录资源来源、建设主体、服务对象、可用状态。")
add_bullet(doc, "资产分类总览：按地区、机构、场景、供给主体、运行状态等维度组织分布视图。")
add_bullet(doc, "资产使用概览：调用频次、覆盖机构、活跃用户、词元消耗、服务热度综合分析，识别高价值/低效/重复资源。")

add_heading(doc, "5.2.2 数据资源 / 知识体系 / 平台工具（保持）", level=3)
add_bullet(doc, "v0.1已有的三类资源子页保留，作为资产总览的下钻视图。")
add_bullet(doc, "【新增】每个资源详情页增加 被以下服务使用 区块，展示资产-服务关联。")

add_heading(doc, "5.2.3 资产-服务关联视图【新增】", level=3)
add_bullet(doc, "以模型/智能体为中心，展示其依赖的数据资源、知识体系、平台工具。")
add_bullet(doc, "以资产为中心，展示使用该资产的所有服务。")
add_bullet(doc, "支撑资源复用决策、重复建设识别、影响范围评估。")

# 5.3 词元计费管理
add_heading(doc, "5.3 词元计费管理", level=2)
add_para(doc, "v0.1将所有子功能混在BillingDashboardPage单页，v0.2拆分为独立子页并补全规则版本化管理。", indent=True)

add_heading(doc, "5.3.1 计费总览", level=3)
add_bullet(doc, "KPI卡片：当前词元总余量、本月已消耗、环比变化、可用服务数量。")
add_bullet(doc, "【新增】KPI数据来源说明：明确每个指标的统计口径、数据源、刷新频率。")
add_bullet(doc, "月度趋势、用量分布、预警记录、实时计费等图表保留。")

add_heading(doc, "5.3.2 计费规则管理（版本化）", level=3)
add_bullet(doc, "三方面差异化能力：分品类差异化计量 / 多主体灵活付费 / 多方分润闭环。")
add_bullet(doc, "【新增】规则版本管理：每个规则有版本号、生效时间、失效时间、状态（草稿/生效/停用）。")
add_bullet(doc, "【新增】规则变更需审批：草稿 → 提交 → 审核 → 生效。")
add_bullet(doc, "【新增】规则与调用明细绑定：每次调用记录使用的规则版本号，便于对账追溯。")
add_bullet(doc, "【新增】规则回滚：发现规则错误可一键回滚至上一版本。")

add_heading(doc, "5.3.3 额度配置管理", level=3)
add_bullet(doc, "按地区、机构、项目、服务和场景分配词元额度。")
add_bullet(doc, "明确服务使用主体、使用范围、可用额度和有效期限。")
add_bullet(doc, "【新增】机构侧配额与平台侧额度配置数据通过事件总线同步，两侧展示一致。")
add_bullet(doc, "【新增】额度调整审批流：发起 → 审核 → 生效，全程留痕。")

add_heading(doc, "5.3.4 消耗监测管理", level=3)
add_bullet(doc, "词元余额变化、阶段消耗、重点服务消耗占比和异常集中调用等情况纳入跟踪监测。")
add_bullet(doc, "【新增】监测告警：余额低于阈值、消耗突增、异常集中调用等触发告警，进入告警工单。")
add_bullet(doc, "【新增】消耗预测：基于历史数据预测剩余可用天数，提前预警。")

add_heading(doc, "5.3.5 词元包管理", level=3)
add_bullet(doc, "普惠包 / 标准包 / 医联体包 / 城市级包等套餐管理。")
add_bullet(doc, "【新增】词元包与机构订阅/配额的关系明确：购买词元包 → 充入机构总账户 → 按配额分配。")
add_bullet(doc, "【新增】词元包到期处理：到期未用完按规则处理（退还/清零/续期）。")

# 5.4 对账查询管理
add_heading(doc, "5.4 对账查询管理", level=2)
add_para(doc, "贯通调用明细、周期账单、异常对账和结果输出。v0.1仅展示，v0.2补全处置流程。", indent=True)

add_heading(doc, "5.4.1 调用明细查询", level=3)
add_bullet(doc, "记录调用时间、调用主体、服务名称、词元消耗、处理状态和异常信息。")
add_bullet(doc, "【新增】调用明细绑定计费规则版本号，便于按规则版本筛选与核对。")
add_bullet(doc, "支撑问题核查、服务评估和责任追溯。")

add_heading(doc, "5.4.2 周期账单管理", level=3)
add_bullet(doc, "按周期归集词元消耗情况，形成地区、机构、项目和服务维度的使用账单。")
add_bullet(doc, "【新增】账单状态机：生成 → 待确认 → 已确认 → 已结算 → 已归档。")
add_bullet(doc, "【新增】机构可在工作台对账单发起异议，进入异常对账流程。")

add_heading(doc, "5.4.3 异常对账管理（含处置）", level=3)
add_bullet(doc, "识别范围：重复调用、异常高频调用、超授权使用、失败扣费、跨机构异常使用等。")
add_bullet(doc, "【新增】异常处置状态机：识别 → 待核验 → 核验中 → 待处置 → 已处置 → 已归档。")
add_bullet(doc, "【新增】处置动作：退费、补扣、追责、转审计。")
add_bullet(doc, "【新增】处置需审批，处置结果同步至机构与开发者。")

add_heading(doc, "5.4.4 对账结果导出", level=3)
add_bullet(doc, "面向内部管理、运营核对、绩效评价和审计留痕需要，提供机构、服务、项目和明细层面的统计结果输出。")
add_bullet(doc, "输出内容保留统计口径、生成时间和必要的追溯信息。")
add_bullet(doc, "【新增】导出操作纳入安全审计，导出文件含水印（操作人+时间）与追溯信息。")
add_bullet(doc, "【新增】导出权限分级：明细级仅平台管理员可导出，机构/服务级运营可导出。")

# 5.5 运营管理
add_heading(doc, "5.5 运营管理", level=2)

add_heading(doc, "5.5.1 服务准入管理", level=3)
add_bullet(doc, "三级差异化审核：高风险（资质核验+技术测评+临床验收）/ 中风险（资质核验+简化技术测试）/ 低风险（备案上架）。")
add_bullet(doc, "【新增】准入审核工作台：待办列表、审核意见、审核记录、退回补正。")
add_bullet(doc, "【新增】准入审核与开发者上架申请双向同步，状态机一致。")

add_heading(doc, "5.5.2 上架下架管理", level=3)
add_bullet(doc, "服务上架、暂停、恢复和退出实行统一管理。")
add_bullet(doc, "明确服务可用范围、适用机构、状态调整原因和处理依据。")
add_bullet(doc, "对资质到期、质量下降或存在安全风险的服务，形成及时处置和退出机制。")
add_bullet(doc, "【新增】下架流程：发起 → 审核 → 公告 → 生效；下架后已订阅机构收到通知与处置方案。")

add_heading(doc, "5.5.3 运行监测管理", level=3)
add_bullet(doc, "综合反映接入机构、上架服务、调用总量、活跃用户、调用趋势和异常预警等情况。")
add_bullet(doc, "【新增】监测大屏支持下钻：从总体指标 → 机构/服务 → 单次调用明细。")
add_bullet(doc, "【新增】告警工单：异常告警自动生成工单，分配处置人，跟踪闭环。")
add_bullet(doc, "【新增】告警规则配置：阈值、触发条件、通知方式、升级策略。")

add_heading(doc, "5.5.4 质量评价管理", level=3)
add_bullet(doc, "AI服务质量评价以准确性、稳定性、响应时效、用户反馈、投诉情况和合规记录等指标为基础。")
add_bullet(doc, "评价结果作为服务推广、额度调整、续用评估和退出管理的重要依据。")
add_bullet(doc, "【新增】评价指标库：维护指标定义、计算口径、数据源、权重。")
add_bullet(doc, "【新增】评价方案管理：按服务类型配置评价方案（指标组合+权重+周期）。")
add_bullet(doc, "【新增】评价结果明细：服务级、机构级、周期级评价结果与趋势。")

add_heading(doc, "5.5.5 安全审计管理", level=3)
add_bullet(doc, "关键业务行为纳入审计管理：服务调用、权限变更、材料审核、额度调整、账单确认等。")
add_bullet(doc, "涉及敏感数据、重要服务和高风险操作的事项纳入重点审计范围。")
add_bullet(doc, "【新增】审计规则配置：定义审计事件、关键字段、风险等级。")
add_bullet(doc, "【新增】审计告警：异常操作（如夜间批量导出、越权尝试）触发告警。")
add_bullet(doc, "【新增】审计日志不可篡改：采用追加写+签名链，支持司法取证。")

# 5.6 运营分析
add_heading(doc, "5.6 运营分析", level=2)
add_para(doc, "独立于运营管理，聚焦数据分析。v0.1三个分析维度在一个页面，v0.2支持各自下钻。", indent=True)

add_heading(doc, "5.6.1 服务热度分析", level=3)
add_bullet(doc, "基于调用次数、覆盖机构、活跃用户、试用转化和词元消耗等指标，形成服务使用热度分析。")
add_bullet(doc, "【新增】支持下钻：服务热度总览 → 单服务热度详情 → 单机构调用明细。")
add_bullet(doc, "【新增】热度趋势对比：支持与上一周期、同类服务对比。")

add_heading(doc, "5.6.2 服务质量分析", level=3)
add_bullet(doc, "基于服务质量评价结果，对AI服务准确性、稳定性、响应时效、用户反馈和合规记录等指标进行综合分析。")
add_bullet(doc, "识别优质服务、问题服务和需整改服务。")
add_bullet(doc, "【新增】质量趋势分析：服务质量随时间变化趋势，识别恶化服务。")
add_bullet(doc, "【新增】质量与热度交叉分析：识别 高热度低质量 服务，优先整改。")

add_heading(doc, "5.6.3 场景应用分析", level=3)
add_bullet(doc, "医保监管、基金管理、临床辅助、疾病筛查、健康服务、科研支撑等领域的AI应用热度和实际需求纳入场景分析。")
add_bullet(doc, "分析结果可作为试点布局、资源投入和场景拓展参考。")
add_bullet(doc, "【新增】场景地图：按地区+场景展示AI应用覆盖情况，识别薄弱场景。")
add_bullet(doc, "【新增】场景供需匹配：对比场景需求与可用服务，识别缺口。")

# 5.7 开发者管理
add_heading(doc, "5.7 开发者管理", level=2)

add_heading(doc, "5.7.1 资质审核管理", level=3)
add_bullet(doc, "开发者资质、产品合规情况、技术能力、数据安全能力和运维保障能力纳入核验范围。")
add_bullet(doc, "涉及医疗器械、算法备案、个人信息处理等要求的开发者，形成审核记录和到期提醒。")
add_bullet(doc, "【新增】资质到期管理：到期前30/15/7天提醒，到期未续证开发者所提供服务自动暂停。")

add_heading(doc, "5.7.2 整改跟踪管理", level=3)
add_bullet(doc, "对测试评估、合规审核和运行监测中发现的问题，统一纳入整改闭环管理。")
add_bullet(doc, "向开发者下达整改要求、责任主体、整改期限并跟踪复核结果。")
add_bullet(doc, "整改情况作为服务上架、续用、暂停合作或清退下线决策的重要依据。")
add_bullet(doc, "【新增】整改工单状态机：下达 → 待整改 → 整改中 → 待复核 → 复核通过/不通过 → 关闭/升级。")
add_bullet(doc, "【新增】整改超期升级：超期未整改自动升级处置（暂停服务/清退）。")

# 5.8 用户与权限管理
add_heading(doc, "5.8 用户与权限管理", level=2)

add_heading(doc, "5.8.1 机构用户管理", level=3)
add_bullet(doc, "医保部门、医疗机构、运营单位和供给单位等不同主体纳入机构用户管理。")
add_bullet(doc, "明确机构属性、业务范围、接入状态和服务授权。")

add_heading(doc, "5.8.2 角色权限管理", level=3)
add_bullet(doc, "构建与医保管理职责相匹配的角色权限体系。")
add_bullet(doc, "区分管理、审核、运营、机构管理、服务使用和开发者等不同权限边界。")
add_bullet(doc, "防止越权查看、越权审核和越权配置。")
add_bullet(doc, "【新增】按3.3权限矩阵实现菜单/按钮/数据三级权限控制。")
add_bullet(doc, "【新增】敏感操作二次确认：导出、下架、删除、额度调整等需二次确认+留痕。")

add_heading(doc, "5.8.3 服务授权管理", level=3)
add_bullet(doc, "服务授权机制覆盖机构、角色、用户、项目和场景，并与词元额度、使用范围和有效期限联动。")
add_bullet(doc, "试点城市、试点医院和专项项目可形成独立授权策略。")

add_heading(doc, "5.8.4 用户行为留痕", level=3)
add_bullet(doc, "用户登录、查询、申请、审核、调用、导出和配置等关键行为纳入留痕管理。")
add_bullet(doc, "额度调整、账单确认、服务下架等重要操作应可追溯到具体责任主体。")

# 5.9 试点与多主体协同管理
add_heading(doc, "5.9 试点与多主体协同管理【新增模块】", level=2)
add_para(doc, "v0.1未实现，v0.2新增。针对PRD 1.2 多主体协同管理 与多次提到的试点场景。", indent=True)

add_heading(doc, "5.9.1 试点项目管理", level=3)
add_bullet(doc, "试点项目定义：试点城市、试点医院、专项项目作为独立授权与计量单元。")
add_bullet(doc, "试点字段：试点名称、试点类型、试点范围（城市/医院/项目）、试点周期、试点政策、独立额度、独立计量规则、成效评估方案。")
add_bullet(doc, "试点状态：筹备 → 启动 → 进行中 → 评估 → 结项 / 推广。")
add_bullet(doc, "试点成效评估：调用情况、覆盖机构、用户反馈、临床/管理成效、成本效益。")

add_heading(doc, "5.9.2 多主体接入边界", level=3)
add_bullet(doc, "明确各类主体（医保部门、医疗机构、运营单位、AI企业、科研机构、医院自研团队）的接入边界。")
add_bullet(doc, "主体字段：主体类型、主体名称、接入角色、可访问模块、数据可见范围、操作权限。")
add_bullet(doc, "【新增】主体协同视图：以服务/项目为中心，展示涉及的多主体与协作关系。")

add_heading(doc, "5.9.3 协同授权与责任留痕", level=3)
add_bullet(doc, "多主体协同场景下的授权规则：主责主体、配合主体、数据共享范围、操作权限。")
add_bullet(doc, "协同操作留痕：每个操作记录主体、角色、时间、对象，支持多主体追溯。")

# 5.10 系统管理
add_heading(doc, "5.10 系统管理", level=2)
add_para(doc, "基础支撑能力，不作为日常业务办理主线模块。", indent=True)

add_heading(doc, "5.10.1 基础参数配置", level=3)
add_bullet(doc, "统一维护平台运行所需的字典、分类、状态、预警阈值、计量参数和通知规则等基础信息。")

add_heading(doc, "5.10.2 接口与集成配置", level=3)
add_bullet(doc, "维护平台与统一身份认证、医保业务系统、数据资源平台、运营服务系统和消息平台等外部系统的基础连接关系。")
add_bullet(doc, "【新增】接口健康检查：定时探活，异常告警。")
add_bullet(doc, "【新增】接口调用统计：调用量、成功率、平均时延。")

add_heading(doc, "5.10.3 API开放平台【新增】", level=3)
add_bullet(doc, "面向开发者提供API开放能力：API文档、API Key管理、调用统计、限流配置。")
add_bullet(doc, "API Key管理：申请、续期、撤销、权限范围（按服务/按机构）。")
add_bullet(doc, "调用统计：调用量、成功率、平均时延、错误分布。")
add_bullet(doc, "限流配置：按Key/按服务/按机构配置QPS与日配额。")

add_heading(doc, "5.10.4 系统日志管理", level=3)
add_bullet(doc, "对登录、配置变更、接口调用、异常告警、数据导出等系统级行为进行记录。")
add_bullet(doc, "为平台运行维护、安全审计和问题排查提供基础依据。")

add_heading(doc, "5.10.5 数据安全配置【新增】", level=3)
add_bullet(doc, "敏感数据识别：标记患者数据、医保数据等敏感字段。")
add_bullet(doc, "数据脱敏规则：按角色配置脱敏策略（如手机号中间四位脱敏）。")
add_bullet(doc, "数据加密：传输加密（HTTPS）、存储加密（敏感字段加密）。")
add_bullet(doc, "权限隔离：按机构/试点项目隔离数据访问范围。")
add_bullet(doc, "数据导出审批：敏感数据导出需审批+留痕。")

doc.add_page_break()

# ===== 6. 非功能性要求 =====
add_heading(doc, "6  非功能性要求", level=1)

add_heading(doc, "6.1 性能要求", level=2)
add_table(doc,
    ["指标", "要求"],
    [
        ["页面首屏加载", "≤ 2秒（常规网络环境）"],
        ["列表查询响应", "≤ 1秒（10万级数据量）"],
        ["大屏数据刷新", "≤ 5秒"],
        ["AI服务调用计量", "实时扣减，延迟 ≤ 1秒"],
        ["并发用户数", "≥ 1000 同时在线"],
        ["日调用处理能力", "≥ 1000万次"],
    ],
    col_widths=[5, 10])

add_heading(doc, "6.2 可用性要求", level=2)
add_bullet(doc, "系统可用性 ≥ 99.9%（年宕机 ≤ 8.76小时）。")
add_bullet(doc, "支持7×24小时运行，计划维护需提前公告。")
add_bullet(doc, "核心模块（计费、对账、调用）支持故障转移，单点故障不影响整体。")

add_heading(doc, "6.3 安全要求", level=2)
add_bullet(doc, "通过等保三级测评。")
add_bullet(doc, "敏感数据传输加密（TLS 1.2+），存储加密（AES-256）。")
add_bullet(doc, "审计日志不可篡改，保留 ≥ 6个月。")
add_bullet(doc, "权限控制遵循最小权限原则，敏感操作二次确认。")
add_bullet(doc, "对接统一身份认证，支持SSO/OAuth。")

add_heading(doc, "6.4 兼容性要求", level=2)
add_bullet(doc, "浏览器：Chrome 100+ / Edge 100+ / Firefox 100+。")
add_bullet(doc, "分辨率：≥ 1280×720，推荐 1920×1080。")
add_bullet(doc, "移动端：响应式适配（机构工作台、门户首页优先）。")

add_heading(doc, "6.5 可维护性要求", level=2)
add_bullet(doc, "前后端分离，接口规范统一（RESTful + OpenAPI 3.0）。")
add_bullet(doc, "配置项与代码分离，运维可调整阈值与规则。")
add_bullet(doc, "日志规范统一，支持日志检索与告警。")

add_heading(doc, "6.6 可扩展性要求", level=2)
add_bullet(doc, "微服务架构，模块可独立部署与扩展。")
add_bullet(doc, "计费规则、评价方案、审计规则等支持热配置。")
add_bullet(doc, "支持横向扩展，应对试点推广后的用户增长。")

doc.add_page_break()

# ===== 7. 端到端业务流程 =====
add_heading(doc, "7  端到端业务流程", level=1)
add_para(doc, "本章描述平台核心端到端业务流程，作为各模块协同的验证视图。", indent=True)

add_heading(doc, "7.1 服务从上架到退出的全生命周期", level=2)
add_para(doc, "流程：开发者上架申请 → 平台准入审核（三级差异化） → 测试协同 → 上架发布 → 机构订阅 → 配额分配 → 调用权限申请 → 服务调用 → 词元计量 → 周期对账 → 质量评价 → 版本迭代/下架退出。", indent=True)
add_note(doc, "关键节点：① 上架申请单通过后生成服务ID与版本ID；② 订阅通过后生成订阅ID与配额ID；③ 调用权限申请通过后生成权限ID；④ 每次调用生成调用ID并绑定计费规则版本；⑤ 周期账单归集调用明细生成账单ID；⑥ 评价与质量评价数据回流至服务档案。", "info")

add_heading(doc, "7.2 词元计量计费闭环", level=2)
add_para(doc, "流程：计费规则制定（版本化） → 额度配置 → 词元包购买/分配 → 服务调用实时扣减 → 调用明细记录 → 周期账单生成 → 机构确认 → 异常对账处置 → 收益分成结算。", indent=True)
add_note(doc, "改进点③：v0.1计费规则无版本管理，导致账单与规则口径不一致风险。v0.2通过规则版本号绑定调用明细与账单，确保对账可追溯。", "warn")

add_heading(doc, "7.3 机构订阅与使用闭环", level=2)
add_para(doc, "流程：机构在服务市场浏览 → 提交订阅申请 → 平台审核 → 订阅生效 → 机构分配院内配额 → 院内用户申请调用权限 → 院内用户调用服务 → 费用结算查询 → 服务评价反馈。", indent=True)
add_note(doc, "改进点④：v0.1机构工作台缺少 调用权限申请 和 评价反馈，导致机构侧闭环断裂。v0.2补全这两个页面，闭环完整。", "warn")

add_heading(doc, "7.4 异常处置闭环", level=2)
add_para(doc, "流程：异常识别（监测/对账） → 异常工单生成 → 待核验 → 核验中 → 待处置 → 处置（退费/补扣/追责/转审计） → 已处置 → 归档。", indent=True)
add_note(doc, "改进点⑤：v0.1异常对账仅识别不处置。v0.2新增完整处置流程与状态机，并支持与安全审计联动。", "warn")

add_heading(doc, "7.5 试点项目闭环", level=2)
add_para(doc, "流程：试点项目立项 → 独立授权 → 独立计量规则 → 试点进行 → 试点数据采集 → 成效评估 → 结项/推广决策。", indent=True)

doc.add_page_break()

# ===== 8. 遗留问题与待确认事项 =====
add_heading(doc, "8  遗留问题与待确认事项", level=1)
add_para(doc, "本章列出本文档尚未明确、需在需求梳理工具中进一步讨论确认的事项。", indent=True)

add_table(doc,
    ["编号", "遗留问题", "建议处理方向", "优先级"],
    [
        ["Q-01", "医保监管角色的具体职责边界与可访问数据范围未明确", "与医保局业务方访谈确认", "高"],
        ["Q-02", "词元与人民币的换算规则、结算周期未明确", "与财务方确认，区分财政统购/套餐预购/专项经费等场景", "高"],
        ["Q-03", "试点项目的独立计量规则与全省统一规则的关系", "建议试点项目可覆盖全省规则，但需独立账本", "高"],
        ["Q-04", "多主体数据共享的合规要求（特别是患者数据）未明确", "与法务/合规方确认，参考《个人信息保护法》《数据安全法》", "高"],
        ["Q-05", "开发者收益分成的具体比例与结算周期未明确", "与运营方确认，按服务类型差异化", "中"],
        ["Q-06", "服务下架后已订阅机构的处置方案未明确（是否退款/如何迁移）", "建议分强制下架与主动下架，差异化处置", "中"],
        ["Q-07", "API开放平台对外提供的接口范围与计费方式未明确", "建议先内部使用，试点稳定后对外开放", "中"],
        ["Q-08", "数据资源/知识体系/平台工具的纳管深度（仅元数据 vs 含实体）未明确", "建议一期仅纳管元数据与访问入口", "中"],
        ["Q-09", "机构工作台是否需要对接院内HIS/HRP系统", "建议一期通过Excel导入/手工维护，二期对接", "低"],
        ["Q-10", "移动端适配范围（仅门户/工作台/管理后台）未明确", "建议一期仅门户与工作台响应式适配", "低"],
        ["Q-11", "审计日志的留存周期与归档策略（6个月是否足够）", "与法务确认，建议 ≥ 1年", "中"],
        ["Q-12", "版本迭代的灰度发布范围与时长（按机构/按比例）", "建议按机构灰度，灰度期1-2周", "低"],
    ],
    col_widths=[1.5, 6, 6, 1.5])

doc.add_page_break()

# ===== 9. 附录 =====
add_heading(doc, "9  附录", level=1)

add_heading(doc, "9.1 v0.1页面与v0.2需求映射", level=2)
add_table(doc,
    ["v0.1页面", "v0.2对应模块", "v0.2状态"],
    [
        ["PortalPage", "4.1 门户首页", "补全公告详情、搜索、个人中心"],
        ["ServiceHallPage", "4.2 AI服务市场", "补全服务详情字段"],
        ["ServiceDetailPage", "4.2.3 服务详情展示", "补全计量规则、试点案例、评价、相关推荐"],
        ["LoginPage", "3.3 角色权限矩阵", "预留SSO/OAuth接口"],
        ["DeveloperOnboardingPage", "4.1.4 开发者入驻", "补全入驻进度查询"],
        ["org/OrgWorkbenchIndexPage", "4.3 机构工作台", "保持"],
        ["org/SubscriptionsPage", "4.3.1 服务订阅管理", "补全退订流程"],
        ["org/QuotaPage", "4.3.2 用量配额管理", "补全调整审批"],
        ["org/BillingQueryPage", "4.3.4 费用结算查询", "补全异常核验"],
        ["—（缺失）", "4.3.3 模型调用权限申请", "新增"],
        ["—（缺失）", "4.3.5 服务评价反馈", "新增"],
        ["—（缺失）", "4.3.6 院内用户管理", "新增"],
        ["DeveloperCenterPage", "4.4 开发者中心", "保持"],
        ["developer/ListingMgmtPage", "4.4.1 上架申请管理", "补全与准入审核双向同步"],
        ["developer/CreateListingPage", "4.4.1 上架申请管理", "保持"],
        ["developer/TestingPage", "4.4.2 测试协同管理", "补全任务状态机"],
        ["—（缺失）", "4.4.3 版本迭代管理", "新增"],
        ["developer/RevenuePage", "4.4.4 收益对账查询", "补全收益争议申诉"],
        ["ModelCatalogPage", "5.1 模型编目管理", "补全资质到期管理"],
        ["DataResourcePage", "5.2 资产总览管理（数据资源）", "纳入资产总览"],
        ["KnowledgeSystemPage", "5.2 资产总览管理（知识体系）", "纳入资产总览"],
        ["PlatformToolsPage", "5.2 资产总览管理（平台工具）", "纳入资产总览"],
        ["—（缺失）", "5.2.1 资产总览首页", "新增"],
        ["—（缺失）", "5.2.3 资产-服务关联视图", "新增"],
        ["BillingDashboardPage", "5.3 词元计费管理", "拆分子页，规则版本化"],
        ["admin/reconciliation/CallDetailsPage", "5.4.1 调用明细查询", "绑定规则版本"],
        ["admin/reconciliation/PeriodBillsPage", "5.4.2 周期账单管理", "补全状态机"],
        ["admin/reconciliation/ExceptionsPage", "5.4.3 异常对账管理", "补全处置流程"],
        ["admin/reconciliation/ExportPage", "5.4.4 对账结果导出", "补全审计与水印"],
        ["admin/operations/AccessReviewPage", "5.5.1 服务准入管理", "补全审核工作台"],
        ["admin/operations/ListingMgmtPage", "5.5.2 上架下架管理", "补全下架流程"],
        ["admin/operations/MonitoringPage", "5.5.3 运行监测管理", "补全下钻与告警工单"],
        ["admin/operations/QualityPage", "5.5.4 质量评价管理", "补全指标库与方案管理"],
        ["admin/operations/AuditPage", "5.5.5 安全审计管理", "补全规则配置与告警"],
        ["OpsAnalyticsPage", "5.6 运营分析", "补全下钻"],
        ["DeveloperMgmtPage", "5.7 开发者管理", "补全资质到期与整改升级"],
        ["UserMgmtPage", "5.8 用户与权限管理", "补全权限矩阵"],
        ["—（缺失）", "5.9 试点与多主体协同管理", "新增模块"],
        ["SettingsPage", "5.10 系统管理", "补全API开放平台与数据安全"],
    ],
    col_widths=[5, 5, 5])

add_heading(doc, "9.2 文档输出说明", level=2)
add_para(doc, "本文档由Claude基于v0.1 Demo源码与PRD评审反馈编制，作为需求梳理工具的输入初稿。文档中标注【新增】【补全】的条目为v0.2相对v0.1的增量。建议在需求梳理工具中按章节逐项确认、细化、补充，形成最终PRD。", indent=True)

# 保存
doc.save(OUT_PATH)
print(f"OK -> {OUT_PATH}")
